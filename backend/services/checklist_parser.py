# backend/services/checklist_parser.py
# Parses engineering checklists from PDF, Excel, CSV, DOCX formats.
# Returns a list of {requirement, expected, notes} dicts.

from __future__ import annotations
import os, io
import pandas as pd

SUPPORTED = {".pdf", ".xlsx", ".xls", ".csv", ".docx", ".doc"}


def parse_checklist(file_path: str) -> list[dict]:
    """
    Parse an engineering checklist file into a list of requirement dicts.
    Each dict: {requirement: str, expected: str, notes: str}
    """
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext in (".xlsx", ".xls"):
        return _parse_excel(file_path)
    elif ext == ".csv":
        return _parse_csv(file_path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported checklist format: {ext}")


def _normalise(rows: list[dict]) -> list[dict]:
    """Ensure every row has requirement, expected, notes keys."""
    out = []
    for r in rows:
        req = (str(r.get("requirement", "") or r.get("Requirement", "") or
               r.get("check", "") or r.get("item", "") or
               list(r.values())[0] if r else "")).strip()
        if not req:
            continue
        exp = str(r.get("expected", "") or r.get("Expected", "") or
                  r.get("status", "") or "Required").strip() or "Required"
        notes = str(r.get("notes", "") or r.get("Notes", "") or
                    r.get("remarks", "") or "").strip()
        out.append({"requirement": req, "expected": exp, "notes": notes})
    return out


def _parse_pdf(path: str) -> list[dict]:
    import fitz
    doc = fitz.open(path)
    rows = []
    for page in doc:
        text = page.get_text()
        for line in text.splitlines():
            line = line.strip()
            if len(line) > 4:
                rows.append({"requirement": line, "expected": "Required", "notes": ""})
    doc.close()
    return rows


def _parse_excel(path: str) -> list[dict]:
    df = pd.read_excel(path, header=0)
    df.columns = [str(c).lower().strip() for c in df.columns]
    return _normalise(df.to_dict("records"))


def _parse_csv(path: str) -> list[dict]:
    df = pd.read_csv(path, header=0)
    df.columns = [str(c).lower().strip() for c in df.columns]
    return _normalise(df.to_dict("records"))


def _parse_docx(path: str) -> list[dict]:
    from docx import Document
    doc = Document(path)
    rows = []
    # Extract from tables first
    for table in doc.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            vals = [c.text.strip() for c in row.cells]
            r = dict(zip(headers, vals))
            rows.append(r)
    # If no tables, fall back to paragraphs
    if not rows:
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                rows.append({"requirement": t, "expected": "Required", "notes": ""})
    return _normalise(rows)
